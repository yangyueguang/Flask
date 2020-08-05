# coding=utf-8
import fitz
import json
import abc
import math
import codecs
import hashlib
import numpy as np
from docx.shared import Pt
from docx.oxml.ns import qn
from docx import Document as DocxDocument


class Dict(dict):
    def __init__(self, sdict={}):
        super(dict, self).__init__()
        if sdict is not None:
            for sk, sv in sdict.items():
                if isinstance(sv, dict):
                    self[sk] = Dict(sv)
                else:
                    self[sk] = sv

    """ A dict that allows for object-like property access syntax. """
    def __getattr__(self, name, default=None):
        try:
            return default if name not in self and default is not None else self[name]
        except KeyError:
            # raise AttributeError(name)
            return None


class Serializable(object):
    __metaclass__ = abc.ABCMeta

    @abc.abstractmethod
    def dump(self):
        pass

    @property
    @abc.abstractmethod
    def text(self):
        pass

    @property
    def type(self):
        return self.__class__.__name__.lower()

    @classmethod
    def load(cls, json_str, **kwargs):
        pass

    @staticmethod
    def md5(string):
        md = hashlib.md5()
        md.update(string.encode())
        return md.hexdigest()

    @staticmethod
    def load_json(json_file):
        with codecs.open(json_file, 'r', encoding='utf-8') as f:
            json_dict = json.load(f, encoding='utf-8')
        return json_dict

    def __str__(self):
        return self.text


class Box(Serializable):
    def __init__(self, chars, x=None, y=None, width=None, height=None, color=0):
        if x is None:
            if chars:
                left = int(math.floor(min([c.x for c in chars])))
                right = int(math.ceil(max([c.x + c.width for c in chars])))
                top = int(math.floor(min([c.y for c in chars])))
                bottom = int(math.ceil(max([c.y + c.height for c in chars])))
                x, y, width, height = left, top, right - left, bottom - top
            else:
                x, y, width, height = 0, 0, 0, 0
        super(Box, self).__init__()
        self.color = color
        self.x = x
        self.y = y
        self.width = width
        self.height = height
        self.chars = chars

    @property
    def text(self):
        return ''.join([c.str for c in self.chars])

    def dump(self):
        return {
            'x': self.x,
            'y': self.y,
            'width': self.width,
            'height': self.height,
            'text': self.text
        }


class BaseElement(Box):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(BaseElement, self).__init__([c for l in lines for c in l.chars], x=x, y=y, width=width, height=height)
        self.lines = lines
        for l in lines:
            l.from_document_element = self
            for c in l.chars:
                c.from_line = l
        self.document = document
        self.parent = None
        self.children = []

    @property
    def global_y(self):
        return self.document.pages[self.chars[0].page - 1].page_info.y + self.y

    def clear_line_ownership(self):
        for line in self.lines:
            line.from_document_element = None

    @classmethod
    def load(cls, json_data, document=None, lines=None):
        x = json_data.get('x', 0)
        y = json_data.get('y', 0)
        width = json_data.get('width', 0)
        height = json_data.get('height', 0)
        element = cls(document, lines if lines else [], x, y, width, height)
        return element

    def dump(self):
        return {
            'x': self.x,
            'y': self.y,
            'width': self.width,
            'height': self.height,
            'type': self.type,
        }

    def __lt__(self, other):
        is_less_than = False
        if self.global_y < other.global_y:
            is_less_than = True
        elif self.global_y == other.global_y:
            if self.x < other.x:
                is_less_than = True
        return is_less_than

    def __str__(self):
        return self.text


class Char(Box):
    __slots__ = ('x', 'y', 'width', 'height', 'font', 'str', 'page', 'index', 'id', 'offset', 'type', 'from_line')

    def __init__(self, x, y, width, height, value='', font=None, page=0, index=0, offset=-1, color=0):
        self.x = x
        self.y = y
        self.width = width
        self.height = height
        self.color = color
        self.font = font
        self.str = value
        self.page = page
        self.index = index
        self.offset = offset
        self.from_line = None
        self.id = str(self.page) + '#' + str(self.offset)

    def dump(self):
        data = {
            'x': self.x,
            'y': self.y,
            'width': self.width,
            'height': self.height,
            'font': self.font,
            'str': self.str,
            'page': self.page,
            'index': self.index,
            'offset': self.offset,
        }
        return data

    @classmethod
    def load(cls, json_char, **kwargs):
        page = json_char.get('page', 0)
        index = json_char.get('index', 0)
        offset = json_char.get('offset', -1)
        x = float(json_char.get('x', 0))
        y = float(json_char.get('y', 0))
        height = float(json_char.get('height', 0))
        width = float(json_char.get('width', 0))
        return cls(x, y, width, height, json_char.get('str', ''), json_char.get('font', ''), page, index, offset)

    def __str__(self):
        return json.dumps(self.dump(), ensure_ascii=False)

    def __eq__(self, other):
        return self.id == other.id

    def __hash__(self):
        return hash(self.id)


class TextLine(Box):
    def __init__(self, boxs, x=None, y=None, width=None, height=None):
        Box.__init__(self, [c for box in boxs for c in box.chars], x, y, width, height)
        self.boxs = boxs

    @property
    def mapper(self):
        return [char.offset for char in self.chars]

    def __eq__(self, other):
        return self.text == other.text

    def __cmp__(self, other):
        if self.y == other.y:
            return self.x - other.x
        return self.y - other.y

    def __lt__(self, other):
        if self.y == other.y:
            return self.x < other.x
        return self.y < other.y

    def __hash__(self):
        return hash(self.text)


class Table(BaseElement):
    def __init__(self, document, lines):
        BaseElement.__init__(self, document, lines, 0, 0, 0, 0)
        self.rows = len(lines)
        self.cols = len(lines[0].boxs)
        self.page_number = 0

    @property
    def cells(self):
        return [cell for row in self.lines for cell in row.boxs]

    @property
    def text_matrix(self):
        text_matrix = [[cell.text for cell in line.boxs] for line in self.lines]
        return np.asarray(text_matrix)


class Title(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(Title, self).__init__(document, lines=lines, x=x, y=y, width=width, height=height)


class Paragraph(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(Paragraph, self).__init__(document, lines=lines, x=x, y=y, width=width, height=height)


class Header(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(Header, self).__init__(document, lines=lines, x=x, y=y, width=width, height=height)


class Footer(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(Footer, self).__init__(document, lines=lines, x=x, y=y, width=width, height=height)


class Catelog(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        super(Catelog, self).__init__(document, lines=lines, x=x, y=y, width=width, height=height)


class Graph(BaseElement):
    def __init__(self, document, lines, x=None, y=None, width=None, height=None):
        BaseElement.__init__(self, document, lines, x, y, width, height)


class Page(BaseElement):
    def __init__(self, document, meta_list, x=None, y=None, width=None, height=None, index=0, rotate=0, areas=None):
        BaseElement.__init__(self, document, meta_list, x, y, width, height)
        self.index = index
        self.rotate = rotate
        self.areas = areas if areas else []

    def dump(self):
        page = {
            'x': self.x,
            'y': self.y,
            'width': self.width,
            'height': self.height,
            'index': self.index,
            'rotate': self.rotate,
            'lines': [tl.dump() for tl in self.lines],
            'areas': [area.dump() for area in self.areas],
        }
        return page

    @classmethod
    def load(cls, json_page, **kwargs):
        info = Dict(json_page)
        lines = [tl.load() for tl in info.lines]
        return cls(None, lines, info.x, info.y, info.width, info.height, info.index, info.rotate, info.areas)

    def __str__(self):
        return json.dumps(self.dump(), ensure_ascii=False)

    def __cmp__(self, other):
        return self.index - other.index

    def __eq__(self, other):
        return self.index == other.index

    def __lt__(self, other):
        return self.index < other.index

    def __hash__(self):
        return hash(id(self))


class Document(Serializable):
    def __init__(self, pages, info=None):
        for page in pages:
            page.document = self
        self.pages = pages
        self.info = Dict(info)

    def dump(self):
        serialized_document = {
            'md5': self.md5,
            'info': self.info,
            'pages': [page.dump() for page in self.pages],
        }
        return serialized_document

    @property
    def text(self):
        pages_text = [p.text for p in self.pages]
        return ''.join(pages_text)

    @classmethod
    def load(cls, json_str, **kwargs):
        page_list = [Page.load(json_page) for json_page in json_str['pages']]
        document = cls(page_list, info=json_str['info'])
        return document

    @classmethod
    def parse(cls, file_name):
        document = fitz.open(file_name)
        pages = []
        index = 0
        global_y = 0
        for page_num in range(document.pageCount):
            page = document.loadPage(page_num)
            page_raw = page.getDisplayList().getTextPage().extractRAWDICT()
            offset = 0
            lines = []
            for block in page_raw['blocks']:
                boxs = []
                for line in block['lines']:
                    text_lines = []
                    for span in line['spans']:
                        cs = []
                        color = span['color']
                        for char in span['chars']:
                            x0, y0, x1, y1 = char['bbox']
                            c = Char(x0, y0, x1 - x0, y1 - y0, char['c'], span['font'], page_num, index, offset, color)
                            index += 1
                            offset += 1
                            cs.append(c)
                        x0, y0, x1, y1 = span['bbox']
                        box = Box(cs, x0, y0, x1 - x0, y1 - y0, span['color'])
                        boxs.append(box)

                    x0, y0, x1, y1 = line['bbox']
                    lin = TextLine(boxs, x0, y0, x1 - x0, y1 - y0)
                    text_lines.append(lin)
                x0, y0, x1, y1 = block['bbox']
                lin = Paragraph(None, text_lines, x0, y0, x1 - x0, y1 - y0)
                lines.append(lin)
            global_y += page_raw['height']
            pages.append(Page(None, lines, 0, global_y, page_raw['width'], page_raw['height'], page.rotation, []))
            info = {
                'md5': Serializable.md5(file_name),
                'name': file_name,
            }
            info.update(document.metadata)
        return Document(pages, info=info)

    # 写入word文本文档
    def save_to_docx(self, name):
        document_element_list = [l for page in self.pages for l in page.lines]
        document_element_list.sort()
        docx_document = DocxDocument()
        docx_document.styles['Normal'].font.name = '宋体'
        docx_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        def add_paragraph(element, docx, add_break=False, bold=False, indent=None):
            paragraph = docx.add_paragraph()
            for line in element.lines:
                font_size = line.chars[0].height
                line_text = line.text
                if add_break:
                    line_text += '\n'
                cur_run = paragraph.add_run(line_text)
                cur_run.font.size = Pt(font_size)
                if bold:
                    cur_run.font.bold = True
            if indent:
                paragraph.paragraph_format.first_line_indent = Pt(2 * font_size)

        def add_table(document_element, docx_document, style='Table Grid', real_column_width=True):
            docx_table = docx_document.add_table(rows=document_element.rows, cols=document_element.cols)
            for row_idx, cell_row in enumerate(document_element.cells):
                for col_idx, box in enumerate(cell_row):
                    docx_table.rows[row_idx].cells[col_idx].text = box.text
                    if real_column_width:
                        docx_table.rows[row_idx].cells[col_idx].width = Pt(box.width)
            docx_table.style = style

        for document_element in document_element_list:
            document_element_class_name = type(document_element).__name__.lower()
            if document_element_class_name == 'Paragraph':
                add_paragraph(document_element, docx_document, indent=True)
            elif document_element_class_name == 'Catelog':
                add_paragraph(document_element, docx_document, add_break=True)
            elif document_element_class_name == 'Title':
                add_paragraph(document_element, docx_document, bold=True)
            elif document_element_class_name == 'Table':
                add_table(document_element, docx_document)
            else:
                pass
        docx_document.save(name)
